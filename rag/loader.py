from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}


@dataclass
class Document:
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentLoader:
    def load(self, path: Path) -> list[Document]:
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            raise ValueError(f"Unsupported document type: {suffix}")
        if suffix in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".pdf":
            text = self._load_pdf(path)
        elif suffix == ".docx":
            text = self._load_docx(path)
        else:
            raise ValueError(f"Unsupported document type: {suffix}")
        return [Document(text=text, metadata={"source": str(path), "suffix": suffix})]

    @staticmethod
    def _load_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("Install pypdf to load PDF files.") from exc

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _load_docx(path: Path) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("Install python-docx to load Word files.") from exc

        document = DocxDocument(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
