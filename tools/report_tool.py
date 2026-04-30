from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from agent.planner import AgentPlan


class ReportTool:
    def __init__(self, reports_dir: Path) -> None:
        self.reports_dir = Path(reports_dir)
        self.reports_dir.mkdir(parents=True, exist_ok=True)

    def generate_business_report(
        self,
        topic: str,
        answer: str,
        plan: AgentPlan,
        sources: list[dict[str, Any]],
        tool_outputs: list[dict[str, Any]],
    ) -> Path:
        slug = self._slugify(topic)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = self.reports_dir / f"{timestamp}_{slug}.md"
        markdown = self._markdown(topic, answer, plan, sources, tool_outputs)
        path.write_text(markdown, encoding="utf-8")
        self._try_write_docx(path.with_suffix(".docx"), topic, markdown)
        return path

    @staticmethod
    def _markdown(
        topic: str,
        answer: str,
        plan: AgentPlan,
        sources: list[dict[str, Any]],
        tool_outputs: list[dict[str, Any]],
    ) -> str:
        source_lines = []
        for idx, source in enumerate(sources, start=1):
            metadata = source.get("metadata", {})
            source_lines.append(
                f"{idx}. {metadata.get('source', 'knowledge_base')} "
                f"(chunk={metadata.get('chunk_id', '-')}, score={source.get('score', 0):.4f})"
            )
        tool_lines = [f"- {item['tool']}: `{item['result']}`" for item in tool_outputs]
        return "\n\n".join(
            [
                f"# {topic}",
                f"生成时间：{datetime.now().isoformat(timespec='seconds')}",
                f"任务类型：`{plan.intent}`",
                "## 执行计划\n" + "\n".join(f"- {step}" for step in plan.steps),
                answer,
                "## 工具调用记录\n" + ("\n".join(tool_lines) if tool_lines else "未调用外部工具。"),
                "## 来源清单\n" + ("\n".join(source_lines) if source_lines else "暂无来源。"),
            ]
        )

    @staticmethod
    def _try_write_docx(path: Path, topic: str, markdown: str) -> None:
        try:
            from docx import Document
        except ImportError:
            return
        document = Document()
        document.add_heading(topic, level=1)
        for block in markdown.split("\n\n"):
            if block.startswith("# "):
                continue
            if block.startswith("## "):
                document.add_heading(block.replace("## ", "", 1), level=2)
            else:
                document.add_paragraph(block)
        document.save(str(path))

    @staticmethod
    def _slugify(value: str) -> str:
        slug = re.sub(r"[^\w\u4e00-\u9fff]+", "-", value.strip().lower()).strip("-")
        return slug[:48] or "business-report"
